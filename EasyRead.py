from tkinter import *
import tkinter as tk
from tkinter import ttk
'''
 *** source code for my "volunteering" software development project
 *** the application was coded using python 3.5.2 - tkinter and docx modules.
 ***  the application basically takes in a docx file and converts it into the desired format which allows people with
 	  <insert here> syndrome to read comfortably.
'''

#importing all the modules needed for the application to run
from tkinter import messagebox
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import math
import random
import time
 
 
'''
	This is a step by step reference for the future me to understand what this code is doing. Hi Future Me! If you really dont understand this, 
	please look into the python docx and tkinter modules documentation for detailed references. I've tried to explain every line of code in detail, 
	so that you can jump right in and make any changes you want.
''' 
main = Tk() #initializing an instance of the tkinter window.
main.title(' ಓದಿ ') #title of the tkinter window. Python converts Kannada directly into unicode
main.iconbitmap('Icon.ico') #this is the icon of the application appearing on the tkinter window`
main.geometry('500x150') #this is the dimensions of the entire window
main.resizable(0,0) # this basically disables the maximize window.

'''
	The next bit of code is something I took right off stack overflow.  i dont exactly know what this does so i'll get back to this. But 
	all i know is it miraculously helps me position my elements in the tkinter notebook frame.
'''
rows = 0
while rows < 50:
    main.rowconfigure(rows, weight=1)
    main.columnconfigure(rows, weight=1)
    rows += 1
 
nb = ttk.Notebook(main) #initializing ttk notebook instance. the ttk notebook is basically a bunch of frames that look like the tabs in any browser.
nb.grid(row=1, column=0, columnspan=50, rowspan=49, sticky='NESW') # defines the layout of the notebook
page1 = ttk.Frame(nb) # this frame is to display the copyright information of the application.
nb.add(page1, text='Copyright') # appending the frame instance to the notebook instance. The title of this frame is 'Copyright'

# string that stores the copyright information
copyright_law = "\n  Copyright ©   2018 by Nishant Ravi Shankar\n     All rights reserved. No part of this software may be reproduced,\n     distributed, or transmitted in any form or by any means,\n     including electronic or mechanical methods, without the prior written\n     permission of the aforementioned copyright holder."

ttk.Label(page1,text = copyright_law,font = "Calibri").place(x = 0,y = 0,width = 1000) #creating a tkinter label for the copyright information and placing it in the frame.

page2 = ttk.Frame(nb) # this frame is to display the "main" frame which actually helps you to upload word files and convert them.
nb.add(page2, text='Convert') # appending the frame instance to the notebook instance. Title of this frame is 'Convert'

# similarly the next two lines of code are for the settinges frame, where you can apply different settings to get a desired output.
page3 = ttk.Frame(nb) 
nb.add(page3, text='Settings')



''' this function is executed when the "browse" button is clicked. the function basically uses the inbuilt tkinter filedialog instance and opens the generic file browsing window.
    the browse window has been specified to specifically search for .docx files, but you can also search for other files (which wont be accepted anyways). Once the file has been 
	selected the returned file path is displayed on an entry widget (using the set method) for the user's reference.
'''
def openFile():
	page2.filename =  filedialog.askopenfilename(initialdir = "/",title = "Select file",filetypes = (("word files","*.docx"),("all files","*.*")))
	entry.set(page2.filename)

'''
	this is the flagship function of the application which does the actual work.
'''
def convertFile():
	settings_file = open("settings.txt","r") #open the settings file that has been bundled along with the application.
	settings_list = settings_file.read().splitlines() #get all the current settings mentioned in the settings file.
	settings = []

	#put the settings into a settings list for future access.
	for setting in settings_list:
		settings.append(setting.split("=")[1])
	
	
	f = entry.get() # get the filepath of the browsed docx file.
	
	#this makes sure that only docx files are uploaded.
	if(f == ""):
		messagebox.showinfo("Failure","Please select a file")
		return 
	if(".docx" not in f):
		messagebox.showinfo("Failure","Invalid File Format")
		return
	if(f.split("/")[-1].split(".")[1] !="docx"):
		messagebox.showinfo("Failure","Invalid File Format")
		return
	
	document = Document(f) #created a docx instance by passing the filepath of the browsed docx file.
	print(f)
	#these three lines of code are to access the name of the docx file
	n = len(f.split("/")) 
	file_path = f.split("/")[1:n-1]
	filename = f.split("/")[-1].split(".")[0]
	filePath = "/"+"/".join(file_path)
	
	paragraphs = [] # list to hold the paragraphs in the file.
	
	# these are the specified [highlight color,font color] combinations
	color_combinations = [["light pink","black"],["light pink","black"],["light pink","black"],["light gray","black"],["light gray","blue"],["light gray","red"],["light yellow","red"],["light yellow","blue"],["light yellow","blue"]]
	
	# these are some of the font colors specified along with their rgb values. Some colors like red and blue required a darker shade and hence the modified rgb values.
	colors = {"red":[180,0,0],"orange":[255,165,0],"pink":[255,192,203],"peach":[255,218,185],"yellow":[255,255,0],"green":[124,252,0],"brown":[139,69,19],"black":[0,0,0],"blue":[0,0,128]}
	
	#["turquoise","orange"],["turquoise","red"],["turquoise","peach"],["turquoise","yellow"],["turquoise","green"]
	
	# these are the highlight colors  used which are specified along with their WD_COLOR_INDEX values. Currently, only a few colors are supported by the library.
	hightlight_color_values = {"light yellow":WD_COLOR_INDEX.YELLOW,"light pink":WD_COLOR_INDEX.PINK,"light gray":WD_COLOR_INDEX.GRAY_25,"turquoise":WD_COLOR_INDEX.TURQUOISE,"light green":WD_COLOR_INDEX.GREEN,"teal":WD_COLOR_INDEX.TEAL}
	
	# this is the zero based indexing list which whose order will be randomized to get a color combination.
	color_combinations_indices = []
	for index in range(0,len(color_combinations)):
		color_combinations_indices.append(index)
	document_1 = Document() # create another docx instance for the formatted file.
	
	# we are now going to loop through the paragraphs of the docx file to extract the text. 'document.paragraphs' returns the list of all the paragraphs in the docx file.
	for p in document.paragraphs:
		paragraphs.append(p.text) # appending the text of the paragraph into the paragraphs list which may help later. Also, when printing the words of the paragraph (for debugging purposes), please encode it in 'utf-8'. this is due to the UnicodeEncodeError
		words = p.text.split(" ")	#splitting the paragraph text into a "list of words".
		paragraph = document_1.add_paragraph() #add_paragraph method adds the paragraph into the new formatted document which returns an instance of the added paragraph which can be used to modify the paragraph next to our needs.
		paragraph_format = paragraph.paragraph_format #paragraph_format method returns an instance whose attributes can be used to format the newly added paragraph text
		paragraph_format.line_spacing = int(settings[2]) #adding line_spacing to the paragraph_format instance. the value is taken from the linespacing value in settings.txt file.
		i = 0
		#Since the color_combinations have to be unpredicatable (to make sure that the student is actually able to read the word by reducing the predicatabity of the color combination order)
		#Hence the color_combinations_indices has to be shuffled.
		random.shuffle(color_combinations_indices) 
		print(color_combinations_indices)
		
		# we need to format each word of the paragraph seperately and hence we're looping throw the "list of words"
		for word in words:
			'''
			The next bunch of lines of code is the CORE IMPLEMENTATION OF THE APPLICATION. As per the requirement, the words need to be formatted-
			the word needs to be highlighted and a font color needs to be added. The highlighting has been done => example: ' Hello ' (where within
			the space within the quotes is highlighted.) You may assume that only the word needs to be highlighted, but 
			an extra space before and after the word are also highlighted so that the word can be read easily. The extra spacing before and after the word
			need to have the same font size and highlight color as the word.This was part of the first update 
			of the application. After every word, spaces have been added to based on the wordspacing value in the bundled settings.txt file.
			'''
			font_size = int(settings[0]) # get the font size for the word
			word_spacing = int(settings[3]) # get the appropriate word spacing from the settings list
			index_length = len(color_combinations_indices) 
			if(word == ''):
				new_word_1 = paragraph.add_run("")
				new_word_1.font_size = Pt(font_size)
				new_word_1.font.highlight_color = hightlight_color_values[color_combinations[color_combinations_indices[i%index_length]][0]]
			else:	
				new_word_1 = paragraph.add_run(" ") #this is the extra spacing before the word begins.
				new_word_1.font.size = Pt(font_size) # set the font size using Pt class from docx.shared module. 
				new_word_1.font.highlight_color = hightlight_color_values[color_combinations[color_combinations_indices[i%index_length]][0]] # set the highlight color from the first index of the element in color_combinations
				
				new_word = paragraph.add_run(word) # this adds the word to the paragraph in the formatted document.
				new_word.font.highlight_color = hightlight_color_values[color_combinations[color_combinations_indices[i%index_length]][0]] # similarly highlighting the word
				new_word.font.name = settings[1] # setting the font name of the word
				font_color = color_combinations[color_combinations_indices[i%index_length]][1] # getting the font color from the second index of the element in color combination
				#print("high = "+hightlight_color_values[color_combinations[color_combinations_indices[i]][0]])
				#print(colors[font_color][0])
				new_word.font.color.rgb = RGBColor(colors[font_color][0],colors[font_color][1],colors[font_color][2]) # sets the font color of the word.
				new_word.font.size = Pt(font_size) # sets the font size of the word
				
				new_word_1 = paragraph.add_run(" ") # this is the extra spacing after the word ends.
				new_word_1.font.size = Pt(font_size) # setting font size of the extra spacing.
				new_word_1.font.highlight_color = hightlight_color_values[color_combinations[color_combinations_indices[i%index_length]][0]] # setting highlight color of the extra spacing.
			
				''' this loop adds the word spacing between the words. it is similar to adding the extra spacing before and after the word, except that it
					has a white background color.
				'''
				for j in range(0,word_spacing):
					new_word_1 = paragraph.add_run(" ")
					new_word_1.font.size = Pt(font_size)
					new_word_1.font.highlight_color = WD_COLOR_INDEX.WHITE
				i+=1
	f = filedialog.asksaveasfilename(defaultextension=".docx") #asks for filePath of new file.
	print(f)
	document_1.save(f) # save it in that filePath
	#print("C:"+filePath+"/"+filename+"_converted"+".docx")
	messagebox.showinfo("Success","File converted") # displays a successful message after the file conversion.



# declaring the widgets and placing them in the 'Convert' frame
ttk.Label(page2,text = "Open",font = "Calibri").place(x = 3,y = 10,width = 300)
entry = StringVar()	# an Entry widget can only be set using the StringVar variable type. 
ttk.Entry(page2,textvariable = entry).place(x = 60,y = 12,width = 350)
ttk.Button(page2,text = "Browse",command = openFile).place(x = 420,y = 8,width = 65)
ttk.Button(page2,text = "Convert",command = convertFile).place(x = 220,y = 95,width = 65)

#declaring the widgets and placing them in the 'Settings' frame
ttk.Label(page3,text = "Font Size",font = "Calibri").place(x = 3,y = 10,width = 300)

#once the application starts, the values have to be taken from the settings.txt file and displayed on the respective widgets as the current values.
settings_file = open("settings.txt","r")
settings_list = settings_file.read().splitlines()
settings = []
for setting in settings_list:
	settings.append(setting.split("=")[1])
print(settings)

# the "changed_<attribute> variables initially have the stored values and they change as soon as someone has applied changes to the settings"
changed_font_size = settings[0]
default = tk.StringVar()
default.set(settings[0])

changed_font_type = settings[1]

changed_line_spacing = settings[2]

changed_word_spacing = settings[3]

#list of valid font sizes taken straight from MS word
font_sizes = ['6','7','8','9','10','10.5','11','12','13','14','15','16','18','20','22','24','26','28','32','36','40','44','48','54','60','66','72','80','88','96']



# this function listens to any changes in the font size made by the user and sets the changed_font_size variable to the new value,
def getFontSize(selection):
	print("Value is  = ")
	global changed_font_size 
	print(changed_font_size)
	changed_font_size = selection
	print(changed_font_size)
font_size_option_menu = tk.OptionMenu(page3,default,*font_sizes,command = getFontSize)
font_size_option_menu.place(x = 90,y = 6,width = 70)


ttk.Label(page3,text = "Font Type",font = "Calibri").place(x = 3,y = 50,width = 300)
font_type = ['Calibri']
default = tk.StringVar()
default.set(font_type[0])





# similarly this function listens to any changes in the font type made by the user and sets the changed_font_size variable to the new value.
def getFontType(selection):
	print("Value is  = ")
	global changed_font_type 
	print(changed_font_type)
	changed_font_type = selection
	
	
	


# the only font type currently available is 'Calibri', the font type option menu is disabled for now.
font_type_option_menu = tk.OptionMenu(page3,default,*font_type,command = getFontType)
font_type_option_menu.place(x = 90,y = 46,width = 100)
#font_type_option_menu.configure(state = "disabled")

ttk.Label(page3,text = "Line Spacing",font = "Calibri").place(x = 250,y = 10,width = 300)


line_spacing_entry = tk.Entry(page3)
line_spacing_entry.place(x = 370,y = 9,width = 50)
line_spacing_entry.insert(0,settings[2])

ttk.Label(page3,text = "Word Spacing",font = "Calibri").place(x = 250,y = 50,width = 300)

word_spacing_entry = tk.Entry(page3)
word_spacing_entry.place(x = 370,y = 49,width = 50)
word_spacing_entry.insert(0,settings[3])


'''this function is executed whenever the Apply button is clicked. It basically asks for a confirmation for changing the values, and if the user says
   ok, it writes the changes made to the bundled settings.txt file. 
'''

def applyChanges():
	apply_changes = messagebox.askokcancel("Message","Apply Changes?")
	if(apply_changes == True):
		changed_line_spacing = line_spacing_entry.get()
		changed_word_spacing = word_spacing_entry.get()
		print("hello world")
		f = open("settings.txt",'w')
		s = "fontsize="+changed_font_size+"\nfonttype="+changed_font_type+"\nlinespacing="+changed_line_spacing+"\nwordspacing="+changed_word_spacing
		f.write(s)


ttk.Button(page3,text = "Apply",command = applyChanges).place(x = 430,y = 90,width = 50) 
main.mainloop()